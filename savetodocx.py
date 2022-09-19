# -*- coding: utf-8 -*-
"""
Created on Tue Apr 20 10:16:19 2021

@author: Souparno
"""

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor



def todocx(search_query,summary):
    document = Document()
    run = document.add_paragraph().add_run()
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(18)
    font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    print('1')
    
    document.add_heading('Medical Writing Automation', 0)
    
    # p = document.add_paragraph('A plain paragraph having some ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    
    document.add_heading('search terms :'+ search_query, level=1)
    document.add_paragraph('Generated Methodology' )
    print('here')
    document.add_paragraph(summary)
    
    
    document.add_page_break()
        
    document.save('D://NLG_TCS_project//selected scripts//'+search_query+' methodology.docx')
    print('doc saved')