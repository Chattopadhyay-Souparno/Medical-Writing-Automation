# -*- coding: utf-8 -*-
"""
Created on Fri Apr 16 10:05:36 2021

@author: Souparno
"""
import torch
import streamlit as st
from methodology_bs4 import *
from contextlib import contextmanager, redirect_stdout
from io import StringIO
from time import sleep
from summarygenerator import SummarizerText
#import textwrap
#from savetodocx import *
from docx import Document
from docx.shared import Inches,Pt,RGBColor
#from docx.shared import Pt
#from docx.shared import RGBColor
from transformers import GPT2Model,GPT2Tokenizer,GPT2Config,BertModel,BertTokenizer,BertConfig,T5Tokenizer, T5ForConditionalGeneration
from summarizer import Summarizer
import requests
import json
from keywordsgenerator import *
import pymongo

client=pymongo.MongoClient('mongodb://127.0.0.1:27017/')
mydb=client['Medical_Writing_Automation']
article=mydb.generatedArticle



st.set_page_config(layout="wide")
@contextmanager
def st_capture(output_func):
    with StringIO() as stdout, redirect_stdout(stdout):
        old_write = stdout.write

        def new_write(string):
            ret = old_write(string)
            output_func(stdout.getvalue())
            return ret
        
        stdout.write = new_write
        yield


@st.cache(allow_output_mutation=True)
def T5_function():  # try to break Streamlit's argument-hashing
    tokenizer = T5Tokenizer.from_pretrained('t5-base',local_files_only=True)
    model = T5ForConditionalGeneration.from_pretrained('t5-base', return_dict=True,local_files_only=True)
    return tokenizer,model # try to break Streamlit's return-value-hashing

@st.cache(allow_output_mutation=True)
def GPT2_function():
    # url = "https://rewriter-paraphraser-text-changer-multi-language.p.rapidapi.com/rewrite"      #APi1
    # headers = {
    # 'content-type': "application/json",
    # 'x-rapidapi-key': "2d8926693fmsh664f75bf9fe13e7p1d82a0jsne01e5012ea3e",
    # 'x-rapidapi-host': "rewriter-paraphraser-text-changer-multi-language.p.rapidapi.com"}

    url = "https://api.promptapi.com/paraphraser"
    
    headers= { "apikey": "arsM5vyqjHuiWofmPZqhdksK0jm2rFP0"    }


    custom_config = GPT2Config.from_pretrained('gpt2/models',local_files_only=True)
    custom_config.output_hidden_states = True
    custom_tokenizer = GPT2Tokenizer.from_pretrained('gpt2/models',local_files_only=True)
    custom_model = GPT2Model.from_pretrained('gpt2/models', config=custom_config)
    model = Summarizer(custom_model=custom_model, custom_tokenizer=custom_tokenizer)
    return url,model,headers

def todocx(search_query,summary):
    document = Document()
    run = document.add_paragraph().add_run()
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(18)
    font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
    print('1')
    
    document.add_heading('Medical Writing Automation', 0)
    
    
    document.add_heading('search terms :'+ search_query, level=1)
    document.add_heading('Generated Methodology' ,2)
    print('here')
    document.add_paragraph(summary)
    
    
    document.add_page_break()
        
    document.save('D://NLG_TCS_project//selected scripts//'+search_query+' methodology.docx')
    print('doc saved')
    return None




# tokenizer = T5Tokenizer.from_pretrained('t5-base',local_files_only=True)
# model = T5ForConditionalGeneration.from_pretrained('t5-base', return_dict=True,local_files_only=True)

tokenizer,model=T5_function()


url,model2,headers=GPT2_function()

st.sidebar.title('Medical Writing Automation')
output= st.empty()
search_query= st.sidebar.text_input("Enter the Medical terms to be searched", )
no_articles_fetched = st.sidebar.number_input('Enter the no. of articles to be fetched',value=1)
#no_articles_fetched=5
section = st.sidebar.selectbox("Article Section to be generated",('','Introduction', 'Methodology', 'Discussion'))
gen_type = st.sidebar.selectbox("Generation type",('','Summarization','Generation'))
if gen_type=='Summarization':
    typ=st.sidebar.selectbox("Summarization type",('','Extraction','Abstraction (GPT2)'))
submit=st.sidebar.button('Generate')







if submit:
    if section=='Methodology':
        if gen_type=='Summarization':
            if typ=='Extraction':
    #            if section=='Methodology':
                    s=st.sidebar.button('Save as Word File')
                    with st_capture(output.code):
                        filename,df,meth_ids=Methodology(search_query,no_articles_fetched)
                        st.write('Methodology Section found in: ',*meth_ids)
                        methodologysummary=SummarizerText(df,5)
                        inputs = tokenizer.encode("summarize: " + methodologysummary,return_tensors='pt',max_length=2048,truncation=True)
                        summary_ids = model.generate(inputs, max_length=200, min_length=30, length_penalty=10.,num_return_sequences=10,early_stopping=False, num_beams=10)                
                        summary = tokenizer.decode(summary_ids[0])                
                        st.header('Methodology: ')
                        text=st.text_area(label='Methodology generated: ',value=summary,height=300)
                        #t=todocx(search_query,text)
                        record={
                                'search terms':search_query,
                                'PMCID': meth_ids,
                                'Article Section':section,
                                'Generation Type':gen_type,
                                'Summarization Type':typ,
                                'Generated Text':summary
                                }
                        article.insert_one(record)
            elif typ=='Abstraction (GPT2)':
                if section=='Methodology':
                    s=st.sidebar.button('Save as Word File')
                    with st_capture(output.code):
                        filename,df,meth_ids=Methodology(search_query,no_articles_fetched)
                        st.write('Methodology Section found in: ',*meth_ids)
                        methodologysummary=SummarizerText(df,5)
                        result = model2(methodologysummary, min_length=30, max_length=1000)
                        summary = "".join(result)
                        
                        # payload = "{\r\"language\": \"en\",\r\"strength\": 3,\r\"text\": \""+summary+"\"\r}"     #API1
                        # response = requests.request("POST", url, data=payload.encode('utf-8'), headers=headers)
                        # resp=response.text
                        payload = summary.encode("utf-8")
                        response = requests.request("POST", url, headers=headers, data = payload)
                        
                        status_code = response.status_code
                        result = response.text
                        dictionary=json.loads(result)
    
                        #st.write(resp['rewrite'])
                        st.header('Methodology: ')
                        text=st.text_area(label='',value=dictionary['paraphrased'],height=300)
                        #t=todocx(search_query,text)                    
                        record={
                                'search terms':search_query,
                                'PMCID': meth_ids,
                                'Article Section':section,
                                'Generation Type':gen_type,
                                'Summarization Type':typ,
                                'Generated Text':dictionary['paraphrased']
                                }
                        article.insert_one(record)
        
        elif gen_type=='Generation':
                s=st.sidebar.button('Save as Word File')
                with st_capture(output.code):
                    filename,df,meth_ids=Methodology(search_query,no_articles_fetched)
                    keyw,medicalkeyw=keywords_gen(df,search_query)
                    keyw=list(map(str,keyw))
                    medicalkeyw=list(map(str,medicalkeyw))
                    keyw=' '.join(keyw)
                    medicalkeyw=' '.join(medicalkeyw)
                    st.header('Keywords: ')
                    st.write(keyw)
                    st.header('Medical Keywords: ')
                    st.write(medicalkeyw)
                
    
    
    
    